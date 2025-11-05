#!/usr/bin/env python3
"""
Resume Updater Script - ENHANCED VERSION WITH CLAUDE API
Intelligently updates your DevOps resume based on job requirements
"""

from docx import Document
from docx.shared import Pt
import os
import sys
import re
from datetime import datetime
from anthropic import Anthropic

class ResumeUpdater:
    def __init__(self, original_resume_path, output_dir):
        self.original_resume_path = original_resume_path
        self.output_dir = output_dir
        self.doc = None

        # Initialize Claude API client - REQUIRED for production use
        self.claude_client = None
        api_key = os.environ.get('ANTHROPIC_API_KEY')
        if api_key:
            try:
                self.claude_client = Anthropic(api_key=api_key)
                print('✅ Claude API initialized successfully')
            except Exception as e:
                error_msg = f'CRITICAL: Claude API initialization failed: {e}'
                print(f'❌ {error_msg}')
                raise RuntimeError(error_msg)
        else:
            print('⚠️  ANTHROPIC_API_KEY not set')
            print('⚠️  Templates will be used where available, but skills without templates will cause errors')
            print('⚠️  Set ANTHROPIC_API_KEY for full functionality')

        self.tech_terms = [
            'AWS', 'Azure', 'GCP', 'Google Cloud',
            'ECS', 'Fargate', 'Lambda', 'EC2', 'S3', 'VPC', 'ELB', 'CloudFormation',
            'Aurora', 'PostgreSQL', 'MySQL', 'DynamoDB', 'MongoDB', 'RDS',
            'Kubernetes', 'Docker', 'OpenShift', 'Helm', 'ArgoCD', 'Kustomize',
            'Jenkins', 'GitLab CI/CD', 'GitHub Actions', 'Tekton', 'Bamboo', 'TeamCity',
            'Terraform', 'Pulumi', 'Ansible', 'Chef', 'Puppet',
            'Prometheus', 'Grafana', 'DataDog', 'Splunk', 'ELK', 'Nagios',
            'Python', 'Bash', 'Shell', 'Groovy', 'Go', 'Ruby', 'Perl',
            'Apache Kafka', 'Kinesis', 'RabbitMQ', 'Redis',
            'Nginx', 'Apache', 'Tomcat', 'JBoss', 'WebSphere',
            'Git', 'GitHub', 'GitLab', 'Bitbucket', 'SVN',
            'JIRA', 'Confluence', 'ServiceNow',
            'Linux', 'RHEL', 'CentOS', 'Ubuntu', 'Debian',
            'Maven', 'Ant', 'Gradle', 'npm', 'yarn',
            'Istio', 'Linkerd', 'Flux', 'Crossplane',
            'SonarQube', 'Trivy', 'Snyk', 'Checkmarx',
            'CI/CD', 'DevOps', 'GitOps', 'DevSecOps', 'SRE',
            'REST API', 'GraphQL', 'gRPC', 'WebSockets',
            'Microservices', 'BFF', 'API Gateway',
            'SAFe Agile', 'Scrum', 'Kanban',
            'SSL', 'TLS', 'mTLS', 'OAuth', 'SAML',
            'CloudWatch', 'Application Insights', 'New Relic',
            'Selenium', 'JUnit', 'TestNG', 'pytest',
            'YAML', 'JSON', 'XML', 'HCL'
        ]
    
    def load_resume(self):
        """Load the original resume"""
        if not os.path.exists(self.original_resume_path):
            print(f'Error: Resume file not found at {self.original_resume_path}')
            sys.exit(1)
        
        self.doc = Document(self.original_resume_path)
        print(f'✓ Loaded resume: {os.path.basename(self.original_resume_path)}')
    
    def extract_all_skills(self, job_description):
        """Extract ALL skills dynamically from job description - COMPREHENSIVE EXTRACTION"""
        try:
            found_skills = set()

            # Comprehensive tech/skill keywords (expanded list)
            known_skills = [
                # Cloud Platforms
                'AWS', 'Azure', 'GCP', 'Google Cloud', 'Cloud', 'Multi-cloud',
                # AWS Services
                'ECS', 'Fargate', 'Lambda', 'EC2', 'S3', 'VPC', 'RDS', 'DynamoDB', 'Aurora',
                'CloudFront', 'Route 53', 'API Gateway', 'SNS', 'SQS', 'CloudWatch', 'CloudTrail',
                'AWS Organizations', 'Service Control Policies', 'SCPs', 'Control Tower',
                'AWS Config', 'Security Hub', 'GuardDuty', 'IAM', 'KMS', 'Secrets Manager',
                'CloudFormation', 'CDK', 'Systems Manager', 'Parameter Store',
                'Elastic Beanstalk', 'EKS', 'EFS', 'Glacier',
                # Container & Orchestration
                'Kubernetes', 'K8s', 'Docker', 'Container', 'Containerd', 'CRI-O',
                'Helm', 'ArgoCD', 'Argo Workflows', 'Flux', 'OpenShift', 'Rancher',
                'Kustomize', 'Istio', 'Linkerd', 'Service Mesh',
                # CI/CD
                'Jenkins', 'GitLab', 'GitLab CI', 'GitHub', 'GitHub Actions', 'Actions',
                'CI/CD', 'Pipeline', 'Pipelines', 'Continuous Integration', 'Continuous Deployment',
                'Tekton', 'Bamboo', 'TeamCity', 'CircleCI', 'Travis CI', 'Azure DevOps',
                'CodePipeline', 'CodeBuild', 'CodeDeploy',
                # IaC & Configuration Management
                'Terraform', 'Pulumi', 'Ansible', 'Infrastructure as Code', 'IaC',
                'Chef', 'Puppet', 'SaltStack', 'Vagrant',
                # Programming & Scripting
                'Python', 'Bash', 'Shell', 'Scripting', 'Script', 'PowerShell',
                'Go', 'Golang', 'Java', 'Node.js', 'JavaScript', 'TypeScript',
                'Ruby', 'Perl', 'Groovy', '.NET', 'C#',
                # Messaging & Streaming
                'Kafka', 'Apache Kafka', 'Kinesis', 'RabbitMQ', 'Redis', 'Messaging',
                'ActiveMQ', 'NATS', 'Pulsar',
                # Monitoring & Observability
                'Prometheus', 'Grafana', 'DataDog', 'Monitoring', 'Observability',
                'ELK', 'Elasticsearch', 'Logstash', 'Kibana', 'Splunk',
                'New Relic', 'AppDynamics', 'Dynatrace', 'Jaeger', 'Zipkin',
                # Security & Compliance
                'Security', 'DevSecOps', 'Security governance', 'Compliance',
                'IQ scripts', 'SIEM', 'Vulnerability scanning', 'Penetration testing',
                'Snyk', 'Aqua', 'Twistlock', 'Falco', 'OPA', 'Open Policy Agent',
                'Vault', 'HashiCorp Vault', 'SSL', 'TLS', 'mTLS', 'OAuth', 'SAML',
                # Databases
                'PostgreSQL', 'MySQL', 'MongoDB', 'Database', 'SQL', 'NoSQL',
                'Cassandra', 'CouchDB', 'MariaDB', 'Oracle', 'MS SQL', 'Redis',
                # Architecture & Patterns
                'API', 'REST API', 'GraphQL', 'gRPC', 'Microservices', 'Monolith',
                'BFF', 'Backend for Frontend', 'Event-driven', 'Serverless',
                'SOA', 'Service-Oriented Architecture',
                # Methodologies
                'SAFe', 'Agile', 'Scrum', 'Kanban', 'DevOps', 'SRE', 'GitOps',
                'Lean', 'ITIL', 'Six Sigma',
                # Version Control & Collaboration
                'Git', 'GitHub', 'GitLab', 'Bitbucket', 'Version control', 'SVN',
                'Jira', 'Confluence', 'ServiceNow', 'PagerDuty',
                # Build Tools
                'Maven', 'Gradle', 'Ant', 'npm', 'yarn', 'pip', 'Webpack',
                # Web Servers & Load Balancers
                'Nginx', 'Apache', 'HAProxy', 'Traefik', 'Envoy',
                # Testing
                'Selenium', 'JUnit', 'TestNG', 'pytest', 'Jest', 'Mocha',
                'Load testing', 'Performance testing', 'Integration testing',
                # Data Formats
                'YAML', 'JSON', 'XML', 'HCL', 'TOML',
                # OS & Platforms
                'Linux', 'Unix', 'RHEL', 'CentOS', 'Ubuntu', 'Debian', 'Windows Server',
                'Windows Administration', 'System Administration', 'Infrastructure Engineer', 'Infrastructure Engineering',
                # Networking
                'DNS', 'Load balancing', 'CDN', 'VPN', 'Firewall', 'WAF',
                'TCP/IP', 'HTTP', 'HTTPS',
                # AI/ML
                'AI/ML', 'Machine Learning', 'AI requirements', 'ML requirements',
                # Other
                'Automation', 'Orchestration', 'Configuration management',
                'Infrastructure monitoring', 'Application performance monitoring',
                'Log aggregation', 'Incident management', 'Change management',
                'Financial Services'
            ]

            job_desc_lower = job_description.lower()

            # 1. Direct keyword matching with word boundaries
            for skill in known_skills:
                # Use word boundary matching to avoid partial matches
                pattern = r'\b' + re.escape(skill.lower()) + r'\b'
                if re.search(pattern, job_desc_lower):
                    found_skills.add(skill)

            # 2. Extract skills from common patterns
            # Pattern: "experience with X" or "knowledge of X"
            patterns = [
                r'(?:experience|expertise|proficiency|knowledge|understanding|skills?|background)[\s\w]*?(?:in|with|of|using)[\s]+([\w\s\-/+\.()]{2,40}?)(?:\s+(?:and|or|,|\.|to|for|in|is|are|will|must|should|including)|\.|$)',
                r'(?:using|utilize|work with|working with|implement|deploy|manage|maintain|configure|design|build|develop)[\s]+([\w\s\-/+\.()]{2,40}?)(?:\s+(?:and|or|,|\.|to|for|in|is|are|will|must|should|including)|\.|$)',
                r'(?:strong|solid|deep|extensive|proven)[\s]+(?:experience|knowledge|understanding|expertise)[\s]+(?:in|with|of)[\s]+([\w\s\-/+\.()]{2,40}?)(?:\s+(?:and|or|,|\.|to|for|in|is|are|will|must|should|including)|\.|$)',
            ]

            for pattern in patterns:
                matches = re.findall(pattern, job_description, re.IGNORECASE)
                for match in matches:
                    cleaned = match.strip().rstrip(',.').strip()
                    # Validate it's a potential skill (check against known skills)
                    if len(cleaned) > 2 and len(cleaned) < 50:
                        # Check if any known skill is in this phrase
                        for known in known_skills:
                            if known.lower() in cleaned.lower():
                                found_skills.add(cleaned)
                                break

            # 3. Extract bulleted or listed items
            bullet_patterns = [
                r'[•\-\*]\s*([A-Za-z0-9\s\-/+\.()]{3,50}?)(?:\n|$)',
                r'^\s*\d+[\.)]\s*([A-Za-z0-9\s\-/+\.()]{3,50}?)(?:\n|$)',
            ]

            for pattern in bullet_patterns:
                matches = re.findall(pattern, job_description, re.MULTILINE | re.IGNORECASE)
                for match in matches:
                    cleaned = match.strip().rstrip(',.').strip()
                    if len(cleaned) > 2 and len(cleaned) < 50:
                        for known in known_skills:
                            if known.lower() in cleaned.lower():
                                found_skills.add(cleaned)
                                break

            # Clean up and deduplicate
            final_skills = set()
            for skill in found_skills:
                cleaned = skill.strip().rstrip(',.')
                if cleaned and len(cleaned) > 1 and not cleaned.lower() in ['a', 'an', 'the', 'and', 'or']:
                    final_skills.add(cleaned)

            # Remove subsets (e.g., if both "AWS" and "AWS Lambda" exist, keep both)
            return sorted(list(final_skills), key=len, reverse=True)

        except Exception as e:
            print(f"Error extracting skills: {e}")
            return []
    
    def find_missing_skills(self, all_skills, resume_text):
        """Find skills NOT in resume"""
        try:
            missing = []
            for skill in all_skills:
                if not re.search(rf'\b{re.escape(skill)}\b', resume_text, re.IGNORECASE):
                    missing.append(skill)
            return missing
        except Exception as e:
            print(f"Error finding missing skills: {e}")
            return []
    
    def generate_bullets_with_claude(self, missing_skills, job_description, resume_context):
        """Use Claude API to generate contextual bullets based on job description - MUST SUCCEED"""
        if not missing_skills:
            return []

        if not self.claude_client:
            raise RuntimeError("CRITICAL: Claude API client not initialized. Set ANTHROPIC_API_KEY environment variable.")

        print(f'\n🤖 Using Claude API to generate bullets for {len(missing_skills)} skills...')

        # Aggressive retry logic - will keep trying until success
        max_retries = 5
        base_wait = 2  # Start with 2 seconds

        for attempt in range(max_retries):
            try:
                # Limit to 10 skills per API call to manage token usage
                skills_to_generate = missing_skills[:10]

                prompt = f"""You are an expert DevOps resume writer. Generate professional resume bullet points for the following skills based on the job requirements.

**Job Description:**
{job_description[:2000]}

**Skills to highlight:**
{', '.join(skills_to_generate)}

**Existing Resume Context (for style matching):**
{resume_context[:1000]}

**Requirements:**
1. Generate EXACTLY ONE bullet point per skill listed above ({len(skills_to_generate)} bullets total)
2. Each bullet should be 15-25 words
3. Use strong action verbs (Implemented, Architected, Deployed, Configured, Managed, etc.)
4. Include specific technical details and measurable impact where possible
5. Match the professional tone of the existing resume
6. DO NOT add '•' or bullet symbols - just the text
7. Focus on DevOps/Cloud/Infrastructure achievements
8. Make it relevant to the job description

**CRITICAL: Output ONLY the bullet points, one per line. No explanations, no headers, no numbering - just {len(skills_to_generate)} plain text bullet points.**"""

                response = self.claude_client.messages.create(
                    model="claude-sonnet-4-5",  # Latest Sonnet 4.5 model
                    max_tokens=2500,
                    temperature=0.7,
                    messages=[{
                        "role": "user",
                        "content": prompt
                    }],
                    timeout=60.0  # Increased to 60 second timeout
                )

                # Extract bullets from response
                content = response.content[0].text.strip()

                if not content:
                    raise ValueError("Claude returned empty response")

                bullets = []

                # Parse bullets - be more lenient with formatting
                for line in content.split('\n'):
                    line = line.strip()
                    # Remove common prefixes
                    for prefix in ['•', '-', '*', '▪', '○', '●', '→']:
                        if line.startswith(prefix):
                            line = line[1:].strip()
                    # Remove numbering like "1.", "2.", etc.
                    import re
                    line = re.sub(r'^\d+[\.)]\s*', '', line)
                    # Remove markdown bold
                    line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)

                    # Accept lines that look like resume bullets (8+ words minimum)
                    if line and len(line.split()) >= 8:
                        bullets.append(line)

                # Validation: Must get at least 50% of requested bullets
                if len(bullets) < len(skills_to_generate) * 0.5:
                    raise ValueError(f"Insufficient bullets: got {len(bullets)}, need at least {len(skills_to_generate) // 2}")

                print(f'  ✓ Claude API SUCCESS: Generated {len(bullets)} bullets for {len(skills_to_generate)} skills')

                # If we got fewer bullets than skills, pad with remaining skills
                if len(bullets) < len(skills_to_generate):
                    print(f'  ⚠️  Got {len(bullets)} bullets for {len(skills_to_generate)} skills, will retry for missing')
                    # This is a partial success - return what we have

                return bullets[:len(skills_to_generate)]

            except Exception as e:
                error_type = type(e).__name__
                wait_time = base_wait * (2 ** attempt)  # Exponential backoff: 2s, 4s, 8s, 16s, 32s

                print(f'  ✗ Claude API error (attempt {attempt + 1}/{max_retries}): {error_type}: {str(e)[:100]}')

                if attempt < max_retries - 1:
                    print(f'  ↳ Retrying in {wait_time} seconds...')
                    import time
                    time.sleep(wait_time)
                else:
                    # This is critical - API MUST work
                    error_msg = f"CRITICAL FAILURE: Claude API failed after {max_retries} attempts. Last error: {error_type}: {e}"
                    print(f'  ❌ {error_msg}')
                    raise RuntimeError(error_msg)

        raise RuntimeError(f"CRITICAL: Failed to generate bullets after {max_retries} attempts")

    def generate_missing_skills_bullets(self, missing_skills, job_description):
        """Generate ONE contextual bullet per missing skill based on job requirements"""
        bullets = []

        if not missing_skills:
            return bullets

        print(f'\n✨ Generating bullets for {len(missing_skills)} missing skills...')
        print(f'   Skills: {", ".join(missing_skills[:10])}{"..." if len(missing_skills) > 10 else ""}')

        # Comprehensive skill templates with multiple variations
        skill_templates = {
            # AWS Services - Organizations & Governance
            'AWS Organizations': 'Implemented AWS Organizations with multi-account structure and Service Control Policies (SCPs) to enforce security governance and compliance across 50+ AWS accounts',
            'Service Control Policies': 'Designed and deployed Service Control Policies (SCPs) for centralized governance, enforcing security baselines and regulatory compliance',
            'SCPs': 'Configured SCPs to implement guardrails for cloud resource usage and enforce organizational security policies',
            'Control Tower': 'Deployed AWS Control Tower for automated multi-account governance and standardized account provisioning',

            # AWS Security
            'AWS Config': 'Configured AWS Config rules with automated remediation for continuous compliance monitoring and infrastructure drift detection',
            'Security Hub': 'Integrated AWS Security Hub for centralized security findings aggregation and automated compliance reporting',
            'AWS Security Hub': 'Deployed AWS Security Hub to aggregate security findings from GuardDuty, Inspector, and Config across multiple accounts',
            'GuardDuty': 'Enabled Amazon GuardDuty for intelligent threat detection and continuous security monitoring',
            'IAM': 'Architected fine-grained IAM policies implementing least-privilege access and role-based access control (RBAC)',
            'AWS IAM': 'Designed comprehensive IAM strategy with automated policy validation and periodic access reviews',
            'KMS': 'Implemented AWS KMS for encryption key management and data protection at rest and in transit',
            'Secrets Manager': 'Deployed AWS Secrets Manager for automated credential rotation and secure application secrets management',

            # AWS Compute & Container
            'ECS': 'Architected ECS Fargate deployments for containerized microservices with auto-scaling and service discovery',
            'Fargate': 'Migrated applications to ECS Fargate for serverless container orchestration, reducing infrastructure overhead by 40%',
            'Lambda': 'Developed serverless architectures using AWS Lambda with event-driven processing and automatic scaling',
            'EC2': 'Managed EC2 fleet with automated patching, right-sizing recommendations, and reserved instance optimization',
            'EKS': 'Deployed production EKS clusters with managed node groups, pod security policies, and cluster autoscaling',

            # AWS Storage & Database
            'S3': 'Implemented S3 lifecycle policies, versioning, and cross-region replication for disaster recovery',
            'Aurora': 'Managed Aurora PostgreSQL clusters with automated backups, read replicas, and performance insights',
            'Aurora PostgreSQL': 'Optimized Aurora PostgreSQL databases for high availability with multi-AZ deployment and automated failover',
            'RDS': 'Administered RDS instances with automated backups, performance tuning, and security group hardening',
            'DynamoDB': 'Designed DynamoDB tables with optimized partition keys, GSI/LSI indexes, and on-demand scaling',
            'EFS': 'Deployed EFS for shared persistent storage across containerized applications with automatic scaling',

            # AWS Networking
            'VPC': 'Architected secure VPC designs with private/public subnets, NAT gateways, and transit gateway connectivity',
            'CloudFront': 'Configured CloudFront distributions with WAF integration, custom SSL certificates, and edge caching',
            'Route 53': 'Managed Route 53 hosted zones with health checks, failover routing, and latency-based routing',
            'API Gateway': 'Built REST and WebSocket APIs using API Gateway with Lambda integration and custom authorizers',

            # AWS Monitoring & Operations
            'CloudWatch': 'Implemented CloudWatch dashboards, custom metrics, and automated alerting for infrastructure and application monitoring',
            'CloudTrail': 'Configured CloudTrail for audit logging, compliance tracking, and security forensics',
            'Systems Manager': 'Utilized AWS Systems Manager for patch management, configuration compliance, and remote command execution',
            'Parameter Store': 'Managed application configurations using Parameter Store with encryption and version control',

            # Cloud Migration
            'Migration': 'Led cloud migration initiatives from on-premises to AWS with minimal downtime using AWS Migration Hub and Server Migration Service',
            'Cloud Migration': 'Executed large-scale cloud migration projects including lift-and-shift and re-platforming strategies',

            # IaC Tools
            'Terraform': 'Developed reusable Terraform modules for infrastructure provisioning with state management and automated planning',
            'CloudFormation': 'Created CloudFormation templates with nested stacks, custom resources, and automated rollback capabilities',
            'CDK': 'Built infrastructure using AWS CDK with TypeScript for type-safe cloud resource definitions',
            'Pulumi': 'Implemented infrastructure as code using Pulumi with Python for multi-cloud deployments',
            'Infrastructure as Code': 'Championed Infrastructure as Code practices using Terraform and GitOps for version-controlled infrastructure',
            'IaC': 'Implemented IaC best practices with modular designs, drift detection, and automated validation',

            # Configuration Management
            'Ansible': 'Automated server configuration and application deployment using Ansible playbooks with role-based organization',
            'Chef': 'Managed configuration drift using Chef recipes and cookbooks for consistent server provisioning',
            'Puppet': 'Implemented Puppet for configuration management across Linux and Windows server fleets',

            # Containers & Orchestration
            'Kubernetes': 'Architected production Kubernetes clusters with Helm charts, RBAC, network policies, and persistent storage',
            'K8s': 'Managed multi-tenant Kubernetes environments with namespace isolation and resource quotas',
            'Docker': 'Containerized applications using Docker with multi-stage builds, optimized image layers, and security scanning',
            'Container': 'Implemented container orchestration strategies with health checks, rolling updates, and auto-healing',
            'Helm': 'Created Helm charts for application deployment with templated configurations and release management',
            'ArgoCD': 'Deployed ArgoCD for GitOps-based continuous delivery with automated sync and rollback capabilities',
            'Argo Workflows': 'Orchestrated complex CI/CD workflows using Argo Workflows for parallel job execution',
            'Flux': 'Implemented Flux for GitOps automation with automated image updates and Helm release management',
            'OpenShift': 'Administered OpenShift container platform with integrated CI/CD and developer self-service',
            'Rancher': 'Managed multi-cluster Kubernetes deployments using Rancher for centralized cluster management',
            'Istio': 'Configured Istio service mesh for traffic management, security, and observability in microservices',
            'Linkerd': 'Implemented Linkerd for lightweight service mesh with automatic mTLS and traffic splitting',

            # CI/CD Tools
            'Jenkins': 'Built Jenkins declarative pipelines with shared libraries, automated testing, and blue-green deployments',
            'GitHub Actions': 'Developed GitHub Actions workflows for automated build, test, and deployment with matrix strategies',
            'GitLab CI': 'Implemented GitLab CI/CD pipelines with dynamic environments, artifact management, and pipeline caching',
            'GitLab': 'Managed GitLab repositories with branch protection, merge request approvals, and CI/CD integration',
            'CI/CD': 'Designed end-to-end CI/CD pipelines with automated testing, security scanning, and progressive delivery',
            'Pipeline': 'Optimized CI/CD pipeline performance reducing build times by 60% through parallelization and caching',
            'Tekton': 'Implemented Tekton pipelines for cloud-native CI/CD with reusable tasks and triggers',
            'Azure DevOps': 'Configured Azure DevOps pipelines with YAML definitions and integration with Azure services',
            'CircleCI': 'Built CircleCI workflows with Docker layer caching and parallel test execution',
            'CodePipeline': 'Designed AWS CodePipeline for automated deployment with multi-stage approval workflows',
            'CodeBuild': 'Configured CodeBuild projects with custom build environments and artifact publishing',
            'CodeDeploy': 'Implemented blue/green deployments using CodeDeploy with automated rollback on failure',

            # Programming & Scripting
            'Python': 'Developed Python automation tools for infrastructure provisioning, log analysis, and API integration',
            'Bash': 'Authored Bash scripts for system administration, backup automation, and deployment orchestration',
            'Shell': 'Created shell scripts for cron jobs, log rotation, and system health monitoring',
            'PowerShell': 'Wrote PowerShell scripts for Windows server automation and Active Directory management',
            'Go': 'Built CLI tools and microservices in Go for high-performance infrastructure automation',
            'Golang': 'Developed concurrent applications in Golang for distributed systems and API services',
            'Groovy': 'Scripted Jenkins shared libraries using Groovy for reusable pipeline components',
            'Node.js': 'Built serverless functions and APIs using Node.js with Express framework',

            # Monitoring & Observability
            'Prometheus': 'Deployed Prometheus for metrics collection with custom exporters and alerting rules',
            'Grafana': 'Created Grafana dashboards for real-time infrastructure and application performance visualization',
            'Loki': 'Implemented Grafana Loki for scalable log aggregation and correlation with metrics and traces',
            'Tempo': 'Deployed Grafana Tempo for distributed tracing with seamless Grafana integration and span storage',
            'DataDog': 'Implemented DataDog for full-stack observability with APM, logs, and infrastructure monitoring',
            'ELK': 'Built centralized logging platform using ELK stack with custom parsers and retention policies',
            'Elasticsearch': 'Managed Elasticsearch clusters for log aggregation and full-text search capabilities',
            'Splunk': 'Configured Splunk for security event correlation and compliance reporting',
            'New Relic': 'Integrated New Relic APM for application performance monitoring and transaction tracing',
            'Jaeger': 'Implemented distributed tracing using Jaeger for microservices troubleshooting',

            # Messaging & Streaming
            'Kafka': 'Architected Kafka clusters for high-throughput event streaming with topic partitioning and replication',
            'Apache Kafka': 'Managed Apache Kafka deployments with Schema Registry and Kafka Connect for data integration',
            'Kinesis': 'Built real-time data pipelines using Kinesis Data Streams with Lambda consumers',
            'RabbitMQ': 'Deployed RabbitMQ for reliable message queuing with high availability and clustering',
            'Redis': 'Implemented Redis for caching, session management, and pub/sub messaging patterns',

            # Databases
            'PostgreSQL': 'Administered PostgreSQL databases with replication, backup strategies, and query optimization',
            'MySQL': 'Managed MySQL databases with master-slave replication and performance tuning',
            'MongoDB': 'Deployed MongoDB replica sets with sharding for horizontal scalability',
            'SQL': 'Optimized complex SQL queries and designed normalized database schemas',
            'NoSQL': 'Architected NoSQL solutions for high-scale applications with eventual consistency patterns',

            # Security & Compliance
            'DevSecOps': 'Integrated security into CI/CD pipelines with SAST/DAST scanning and vulnerability management',
            'Security governance': 'Established security governance frameworks with automated compliance checks and audit trails',
            'Compliance': 'Ensured SOC2, HIPAA, and PCI-DSS compliance through automated controls and regular audits',
            'IQ scripts': 'Developed IQ scripts for custom security policy validation and automated compliance reporting',
            'Snyk': 'Integrated Snyk for container and dependency vulnerability scanning in CI/CD pipelines',
            'Vault': 'Deployed HashiCorp Vault for secrets management with dynamic credentials and encryption as a service',
            'HashiCorp Vault': 'Implemented Vault for centralized secrets storage with automated secret rotation',
            'OPA': 'Configured Open Policy Agent for policy-based access control in Kubernetes',

            # Architectures & Patterns
            'Microservices': 'Designed microservices architecture with service discovery, circuit breakers, and event-driven communication',
            'API': 'Built RESTful APIs with OpenAPI specifications, rate limiting, and OAuth2 authentication',
            'REST API': 'Developed REST APIs following best practices with versioning and comprehensive documentation',
            'GraphQL': 'Implemented GraphQL APIs with schema stitching and efficient data fetching',
            'Serverless': 'Architected serverless applications using Lambda, API Gateway, and DynamoDB for auto-scaling',
            'Event-driven': 'Designed event-driven architectures using SNS/SQS with loose coupling and async processing',
            'BFF': 'Implemented Backend for Frontend pattern for optimized mobile and web API experiences',

            # Methodologies
            'Agile': 'Led Agile development practices with sprint planning, daily standups, and retrospectives',
            'SAFe': 'Practiced SAFe Agile framework for large-scale program coordination and release planning',
            'Scrum': 'Facilitated Scrum ceremonies as Scrum Master ensuring team velocity and continuous improvement',
            'DevOps': 'Championed DevOps culture fostering collaboration between development and operations teams',
            'GitOps': 'Implemented GitOps principles using Git as single source of truth for declarative infrastructure',
            'SRE': 'Applied Site Reliability Engineering practices with SLOs, error budgets, and blameless postmortems',

            # Version Control & Collaboration
            'Git': 'Managed Git workflows with branching strategies, pull request reviews, and merge conflict resolution',
            'GitHub': 'Administered GitHub organizations with branch protection rules and required status checks',
            'Jira': 'Tracked project delivery using Jira with custom workflows and automated reporting',
            'Confluence': 'Maintained technical documentation in Confluence for runbooks and architecture decisions',
            'ServiceNow': 'Managed incident and change requests through ServiceNow with ITIL compliance',

            # Web Servers & Load Balancing
            'Nginx': 'Configured Nginx as reverse proxy and load balancer with SSL termination and rate limiting',
            'Apache': 'Administered Apache web servers with virtual hosts, mod_rewrite rules, and security hardening',
            'HAProxy': 'Deployed HAProxy for high-availability load balancing with health checks and session persistence',

            # Testing
            'Selenium': 'Automated browser testing using Selenium WebDriver for end-to-end UI validation',
            'pytest': 'Wrote comprehensive unit and integration tests using pytest with fixtures and mocking',
            'Load testing': 'Performed load testing using JMeter to identify performance bottlenecks and capacity limits',

            # Networking
            'DNS': 'Managed DNS infrastructure with DNSSEC, zone transfers, and GeoDNS for global traffic distribution',
            'Load balancing': 'Implemented application load balancing with health checks and sticky sessions',
            'VPN': 'Configured site-to-site and client VPN solutions for secure remote access',
            'SSL': 'Managed SSL/TLS certificates with automated renewal using Let\'s Encrypt and ACM',

            # Operating Systems
            'Linux': 'Administered Linux servers (RHEL, Ubuntu) with kernel tuning, security patching, and troubleshooting',
            'RHEL': 'Managed Red Hat Enterprise Linux systems with subscription management and satellite integration',
            'Ubuntu': 'Deployed Ubuntu server infrastructure with unattended upgrades and landscape management',

            # Cost Optimization
            'AWS cost optimization': 'Reduced AWS costs by 35% through reserved instances, spot instances, and resource right-sizing',
            'cost optimization': 'Implemented FinOps practices with cost allocation tags, budget alerts, and optimization recommendations',

            # Additional
            'Automation': 'Automated repetitive operational tasks reducing manual effort by 70% and improving reliability',
            'Orchestration': 'Orchestrated complex multi-tier application deployments with zero-downtime strategies',
        }

        try:
            # Limit to maximum of 15 bullets to avoid overwhelming the resume
            skills_to_process = missing_skills[:15]

            # HYBRID APPROACH: Separate skills into template-matched and unmatched
            template_matched_skills = []
            unmatched_skills = []
            template_bullets = []

            print('\n📊 Categorizing skills:')
            for skill in skills_to_process:
                skill_lower = skill.lower()
                matched = False

                # Try exact match first
                for template_skill, template_bullet in skill_templates.items():
                    if skill_lower == template_skill.lower():
                        template_bullets.append(f'•   {template_bullet}')
                        template_matched_skills.append(skill)
                        matched = True
                        break

                # Try partial match if no exact match
                if not matched:
                    for template_skill, template_bullet in skill_templates.items():
                        if skill_lower in template_skill.lower() or template_skill.lower() in skill_lower:
                            template_bullets.append(f'•   {template_bullet}')
                            template_matched_skills.append(skill)
                            matched = True
                            break

                if not matched:
                    unmatched_skills.append(skill)

            print(f'   ✓ {len(template_matched_skills)} skills matched templates')
            print(f'   ✓ {len(unmatched_skills)} skills need API generation')

            # For unmatched skills, REQUIRE Claude API (no generic fallbacks)
            api_bullets = []
            if unmatched_skills:
                if not self.claude_client:
                    raise RuntimeError(f"CRITICAL: {len(unmatched_skills)} skills have no templates and Claude API is not available. Set ANTHROPIC_API_KEY. Skills: {', '.join(unmatched_skills)}")

                print(f'\n🤖 Using Claude API for {len(unmatched_skills)} unmatched skills: {", ".join(unmatched_skills)}')
                resume_text = '\n'.join([p.text for p in self.doc.paragraphs])

                # This will raise RuntimeError if it fails after retries
                api_bullets_raw = self.generate_bullets_with_claude(unmatched_skills, job_description, resume_text)

                # Add bullet prefix if not present
                for bullet in api_bullets_raw:
                    if not bullet.startswith('•'):
                        api_bullets.append(f'•   {bullet}')
                    else:
                        api_bullets.append(bullet)

                print(f'   ✓ Claude API provided {len(api_bullets)} bullets')

            # Combine template bullets and API bullets
            bullets = template_bullets + api_bullets

            print(f'\n✅ Total bullets generated: {len(bullets)}')
            print(f'   • Template-based: {len(template_bullets)}')
            print(f'   • API/Generic: {len(api_bullets)}')

            return bullets

        except Exception as e:
            print(f"Error generating bullets: {e}")
            import traceback
            traceback.print_exc()
            return []
    
    def parse_requirements(self, requirements_text):
        """Parse job requirements"""
        print('\n📋 Analyzing requirements...')
        
        requirements = {
            'cloud_services': [],
            'containers': [],
            'cicd_tools': [],
            'programming': [],
            'databases': [],
            'monitoring': [],
            'messaging': [],
            'other_skills': [],
            'methodologies': [],
            'missing_skills': [],
            'all_extracted_skills': []
        }
        
        try:
            all_extracted_skills = self.extract_all_skills(requirements_text)
            requirements['all_extracted_skills'] = all_extracted_skills
            print(f'✓ Found {len(all_extracted_skills)} total skills')
            
            resume_text = '\n'.join([p.text for p in self.doc.paragraphs])
            missing_skills = self.find_missing_skills(all_extracted_skills, resume_text)
            requirements['missing_skills'] = missing_skills[:20]
            print(f'✓ Missing: {len(missing_skills)} skills')
            
            req_lower = requirements_text.lower()
            
            if 'ecs' in req_lower or 'fargate' in req_lower:
                requirements['cloud_services'].append('ECS Fargate')
            if 'lambda' in req_lower or 'serverless' in req_lower:
                requirements['cloud_services'].append('Lambda')
            if 'aurora' in req_lower:
                requirements['cloud_services'].append('Aurora PostgreSQL')
            if 'dynamodb' in req_lower:
                requirements['cloud_services'].append('DynamoDB')
            if 'kinesis' in req_lower:
                requirements['cloud_services'].append('Kinesis')
            
            if 'kubernetes' in req_lower or 'k8s' in req_lower:
                requirements['containers'].append('Kubernetes')
            if 'docker' in req_lower:
                requirements['containers'].append('Docker')
            if 'helm' in req_lower:
                requirements['containers'].append('Helm')
            if 'argocd' in req_lower:
                requirements['containers'].append('ArgoCD')
            
            if 'jenkins' in req_lower:
                requirements['cicd_tools'].append('Jenkins')
            if 'github actions' in req_lower:
                requirements['cicd_tools'].append('GitHub Actions')
            if 'gitlab' in req_lower:
                requirements['cicd_tools'].append('GitLab CI/CD')
            
            if 'postgres' in req_lower:
                requirements['databases'].append('PostgreSQL')
            if 'mysql' in req_lower:
                requirements['databases'].append('MySQL')
            if 'mongodb' in req_lower:
                requirements['databases'].append('MongoDB')
            
            if 'kafka' in req_lower:
                requirements['messaging'].append('Apache Kafka')
            
            if 'prometheus' in req_lower:
                requirements['monitoring'].append('Prometheus')
            if 'grafana' in req_lower:
                requirements['monitoring'].append('Grafana')
            
            if 'microservices' in req_lower:
                requirements['other_skills'].append('microservices')
            if 'bff' in req_lower:
                requirements['other_skills'].append('BFF')
            
            if 'safe' in req_lower:
                requirements['methodologies'].append('SAFe Agile')
            
            return requirements
        
        except Exception as e:
            print(f"Error parsing requirements: {e}")
            return requirements
    
    def make_selective_bold(self, paragraph, tech_list):
        """Make tech terms bold"""
        try:
            full_text = paragraph.text
            for run in paragraph.runs:
                run.text = ''
            
            sorted_tech = sorted(tech_list, key=len, reverse=True)
            remaining_text = full_text
            
            while remaining_text:
                earliest_pos = len(remaining_text)
                earliest_term = None
                
                for term in sorted_tech:
                    pos = remaining_text.find(term)
                    if pos != -1 and pos < earliest_pos:
                        earliest_pos = pos
                        earliest_term = term
                
                if earliest_term:
                    if earliest_pos > 0:
                        run = paragraph.add_run(remaining_text[:earliest_pos])
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        run.font.bold = False
                    
                    tech_run = paragraph.add_run(earliest_term)
                    tech_run.font.name = 'Times New Roman'
                    tech_run.font.size = Pt(11)
                    tech_run.font.bold = True
                    
                    remaining_text = remaining_text[earliest_pos + len(earliest_term):]
                else:
                    if remaining_text:
                        run = paragraph.add_run(remaining_text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        run.font.bold = False
                    break
        except Exception as e:
            print(f"Error making bold: {e}")
    
    def insert_summary_bullets(self, bullets):
        """Insert bullets into summary with preserved formatting"""
        if not bullets:
            print('\n✏️  No summary bullets to add')
            return False

        print(f'\n✏️  Adding {len(bullets)} bullets to Summary...')

        try:
            insertion_index = None
            reference_para = None

            # Find BACKGROUND Summary or PROFESSIONAL SUMMARY section
            for i, para in enumerate(self.doc.paragraphs):
                para_text = para.text.strip().upper()

                if 'BACKGROUND' in para_text and 'SUMMARY' in para_text:
                    print(f'  ✓ Found BACKGROUND Summary at line {i}')
                    # Find the last paragraph before next section (TECHNICAL SKILLS or WORK EXPERIENCE)
                    for j in range(i+1, min(i+40, len(self.doc.paragraphs))):
                        next_text = self.doc.paragraphs[j].text.strip().upper()
                        # Stop if we hit another section
                        if 'TECHNICAL SKILLS' in next_text or 'WORK EXPERIENCE' in next_text or 'EDUCATION' in next_text:
                            # Insert before this section
                            reference_para = self.doc.paragraphs[j-1]
                            insertion_index = j
                            print(f'  ✓ Will insert at line {insertion_index} (before {next_text[:30]}...)')
                            break
                    if insertion_index:
                        break
                elif 'PROFESSIONAL SUMMARY' in para_text or para_text == 'SUMMARY':
                    print(f'  ✓ Found summary section at line {i}')
                    # Same logic as above
                    for j in range(i+1, min(i+40, len(self.doc.paragraphs))):
                        next_text = self.doc.paragraphs[j].text.strip().upper()
                        if 'TECHNICAL SKILLS' in next_text or 'WORK EXPERIENCE' in next_text or 'EDUCATION' in next_text:
                            reference_para = self.doc.paragraphs[j-1]
                            insertion_index = j
                            print(f'  ✓ Will insert at line {insertion_index}')
                            break
                    if insertion_index:
                        break

            if insertion_index and reference_para:
                # Copy formatting from reference paragraph (List Paragraph style)
                for bullet_text in reversed(bullets):
                    # Remove bullet character if present (we'll use Word's list formatting)
                    if bullet_text.startswith('•'):
                        bullet_text = bullet_text.lstrip('•').strip()

                    new_para = reference_para.insert_paragraph_before()
                    new_para.text = bullet_text

                    # Preserve ALL formatting attributes including list style
                    new_para.paragraph_format.left_indent = reference_para.paragraph_format.left_indent
                    new_para.paragraph_format.first_line_indent = reference_para.paragraph_format.first_line_indent
                    new_para.paragraph_format.space_before = reference_para.paragraph_format.space_before
                    new_para.paragraph_format.space_after = reference_para.paragraph_format.space_after
                    new_para.paragraph_format.line_spacing = reference_para.paragraph_format.line_spacing
                    new_para.paragraph_format.alignment = reference_para.paragraph_format.alignment

                    # CRITICAL: Copy the List Paragraph style
                    new_para.style = reference_para.style

                    # Copy numbering format if present
                    try:
                        if hasattr(reference_para._element, 'pPr') and reference_para._element.pPr is not None:
                            ref_pPr = reference_para._element.pPr
                            ref_numPr = ref_pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                            if ref_numPr is not None:
                                from copy import deepcopy
                                new_pPr = new_para._element.get_or_add_pPr()
                                new_pPr.append(deepcopy(ref_numPr))
                    except Exception as e:
                        print(f'    Warning: Could not copy numbering: {e}')

                    # Apply selective bold formatting
                    self.make_selective_bold(new_para, self.tech_terms)

                print('  ✓ Summary updated successfully')
                return True
            else:
                print('  ✗ Could not find suitable insertion point in Summary section')
                return False
        except Exception as e:
            print(f"Error inserting summary bullets: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def insert_job_bullets(self, bullets, company_keyword, year):
        """Insert bullets into job section with preserved formatting"""
        if not bullets:
            print(f'\n✏️  No bullets to add to {company_keyword}')
            return False

        print(f'\n✏️  Adding {len(bullets)} bullets to {company_keyword}...')

        try:
            # Find the company section
            company_found = False
            reference_para = None

            for i, para in enumerate(self.doc.paragraphs):
                # Look for company heading (flexible matching)
                if company_keyword.lower() in para.text.lower() or year in para.text:
                    company_found = True
                    # Now find bullet points in this job section
                    for j in range(i, min(i+60, len(self.doc.paragraphs))):
                        para_text = self.doc.paragraphs[j].text.strip()

                        # Look for List Paragraph style (bullets)
                        if (self.doc.paragraphs[j].style.name == 'List Paragraph' and len(para_text) > 10):
                            reference_para = self.doc.paragraphs[j]

                            # Check if this is a good insertion point (after first few bullets)
                            # Look for specific anchors or just use first bullet
                            anchor_keywords = [
                                'Tekton', 'ArgoCD', 'Kubernetes', 'Flux', 'GitOps',
                                'pipeline', 'deployment', 'infrastructure'
                            ]

                            is_anchor = any(keyword.lower() in para_text.lower() for keyword in anchor_keywords)

                            # Insert after the first bullet or after anchor bullet
                            if is_anchor or (j - i) <= 3:  # First few bullets
                                insert_at = j + 1

                                # Get reference for formatting
                                ref_para = self.doc.paragraphs[insert_at] if insert_at < len(self.doc.paragraphs) else reference_para

                                # Insert bullets in reverse order
                                for addition in reversed(bullets):
                                    # Remove bullet character if present
                                    if addition.startswith('•'):
                                        addition = addition.lstrip('•').strip()

                                    new_p = ref_para.insert_paragraph_before()
                                    new_p.text = addition

                                    # Preserve ALL formatting
                                    new_p.paragraph_format.left_indent = reference_para.paragraph_format.left_indent
                                    new_p.paragraph_format.first_line_indent = reference_para.paragraph_format.first_line_indent
                                    new_p.paragraph_format.space_before = reference_para.paragraph_format.space_before
                                    new_p.paragraph_format.space_after = reference_para.paragraph_format.space_after
                                    new_p.paragraph_format.line_spacing = reference_para.paragraph_format.line_spacing
                                    new_p.paragraph_format.alignment = reference_para.paragraph_format.alignment
                                    new_p.style = reference_para.style

                                    # Copy numbering format
                                    try:
                                        if hasattr(reference_para._element, 'pPr') and reference_para._element.pPr is not None:
                                            ref_pPr = reference_para._element.pPr
                                            ref_numPr = ref_pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                                            if ref_numPr is not None:
                                                from copy import deepcopy
                                                new_pPr = new_p._element.get_or_add_pPr()
                                                new_pPr.append(deepcopy(ref_numPr))
                                    except:
                                        pass

                                    # Apply selective bold
                                    self.make_selective_bold(new_p, self.tech_terms)

                                print(f'  ✓ {company_keyword} updated successfully')
                                return True

                        # Stop if we hit next job section (look for years or next company)
                        if j > i and (re.search(r'\b(19|20)\d{2}\b', para_text) or
                                      ('LLC' in para_text or 'Inc' in para_text or 'Corp' in para_text)):
                            # Reached next section, use last found bullet as reference
                            if reference_para:
                                insert_at = j

                                for addition in reversed(bullets):
                                    new_p = reference_para.insert_paragraph_before()
                                    new_p.text = addition

                                    # Preserve formatting
                                    new_p.paragraph_format.left_indent = reference_para.paragraph_format.left_indent
                                    new_p.paragraph_format.first_line_indent = reference_para.paragraph_format.first_line_indent
                                    new_p.paragraph_format.space_before = reference_para.paragraph_format.space_before
                                    new_p.paragraph_format.space_after = reference_para.paragraph_format.space_after
                                    new_p.paragraph_format.line_spacing = reference_para.paragraph_format.line_spacing
                                    new_p.paragraph_format.alignment = reference_para.paragraph_format.alignment
                                    new_p.style = reference_para.style

                                    self.make_selective_bold(new_p, self.tech_terms)

                                print(f'  ✓ {company_keyword} updated successfully')
                                return True
                            break
                    break

            if not company_found:
                print(f'  ✗ Could not find company "{company_keyword}"')
                return False
            elif not reference_para:
                print(f'  ✗ Could not find bullet points in {company_keyword} section')
                return False
            else:
                print(f'  ✗ Could not find suitable insertion point')
                return False

        except Exception as e:
            print(f"Error inserting job bullets: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def update_technical_skills(self, requirements):
        """Update technical skills table"""
        print('\n✏️  Updating Technical Skills...')
        
        try:
            updates_made = 0
            
            for table in self.doc.tables:
                for row in table.rows:
                    cells = row.cells
                    if len(cells) >= 2:
                        category = cells[0].text.strip()
                        content = cells[1].text.strip()
                        
                        if 'Cloud Technologies' in category and requirements['cloud_services']:
                            new_services = [svc for svc in requirements['cloud_services'] if svc not in content]
                            if new_services and 'Amazon Web Services' in content:
                                cells[1].text = content + ', ' + ', '.join(new_services)
                                updates_made += 1
                        
                        elif 'CI/CD Tools' in category and requirements['cicd_tools']:
                            new_tools = [tool for tool in requirements['cicd_tools'] if tool not in content]
                            if new_tools:
                                cells[1].text = content + ', ' + ', '.join(new_tools)
                                updates_made += 1
                        
                        elif 'Databases' in category and requirements['databases']:
                            new_dbs = [db for db in requirements['databases'] if db not in content]
                            if new_dbs:
                                cells[1].text = content + ', ' + ', '.join(new_dbs)
                                updates_made += 1
            
            print(f'  ✓ Technical Skills updated ({updates_made} changes)')
        except Exception as e:
            print(f"Error updating skills: {e}")
    
    def save_resume(self, output_path=None):
        """Save resume"""
        if output_path is None:
            base_name = os.path.splitext(os.path.basename(self.original_resume_path))[0]
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'{base_name}_Updated_{timestamp}.docx'
            output_path = os.path.join(self.output_dir, filename)
        
        os.makedirs(self.output_dir, exist_ok=True)
        self.doc.save(output_path)
        print(f'\n✅ Saved: {os.path.basename(output_path)}')
        return output_path
    
    def verify_keywords_added(self, original_text, updated_text, missing_skills):
        """Verify which keywords were successfully added (no extra API calls)"""
        original_lower = original_text.lower()
        updated_lower = updated_text.lower()

        added_keywords = []
        already_present = []
        not_added = []

        for skill in missing_skills:
            skill_lower = skill.lower()
            was_present = skill_lower in original_lower
            is_present = skill_lower in updated_lower

            if is_present and not was_present:
                added_keywords.append(skill)
            elif is_present and was_present:
                already_present.append(skill)
            else:
                not_added.append(skill)

        return {
            'added': added_keywords,
            'already_present': already_present,
            'not_added': not_added
        }

    def update_resume(self, requirements_text):
        """Main update method - returns (output_path, added_keywords)"""
        print('\n' + '='*60)
        print('RESUME UPDATER - GUARANTEED SKILL COVERAGE')
        print('='*60)

        try:
            self.load_resume()

            # Store original text for verification
            original_text = '\n'.join([p.text for p in self.doc.paragraphs])

            requirements = self.parse_requirements(requirements_text)

            if not any([requirements.get('missing_skills'), requirements.get('cloud_services')]):
                print('\n⚠️  No relevant skills found')
                return None, []

            # Track missing skills for verification
            missing_skills = requirements.get('missing_skills', [])

            # Generate bullets for missing skills
            generated_bullets = []
            if missing_skills:
                generated_bullets = self.generate_missing_skills_bullets(
                    missing_skills,
                    requirements_text
                )

            # Split bullets: 40% to summary, 60% to most recent job
            if generated_bullets:
                split_point = max(1, len(generated_bullets) * 4 // 10)  # 40%
                summary_bullets = generated_bullets[:split_point]
                job_bullets = generated_bullets[split_point:]

                print(f'\n📊 Distributing {len(generated_bullets)} bullets:')
                print(f'   • Summary: {len(summary_bullets)} bullets')
                print(f'   • Job Experience: {len(job_bullets)} bullets')
            else:
                summary_bullets = []
                job_bullets = []

            # Insert bullets
            self.insert_summary_bullets(summary_bullets)

            if job_bullets:
                # Try to find the most recent job (look for 2024, then 2023, etc.)
                job_inserted = False
                for year in ['2024', '2023', '2022', '2021', '2020']:
                    for para in self.doc.paragraphs:
                        if year in para.text and ('Engineer' in para.text or 'Developer' in para.text or 'Architect' in para.text):
                            # Extract company name from paragraph
                            company_name = para.text.split(',')[0].split('-')[0].strip()
                            if self.insert_job_bullets(job_bullets, company_name, year):
                                job_inserted = True
                                break
                    if job_inserted:
                        break

                if not job_inserted:
                    print('  ⚠️  Could not find recent job section, skipping job bullets')
                    # Add them to summary instead
                    print('  ↳ Adding remaining bullets to summary')
                    self.insert_summary_bullets(job_bullets)

            self.update_technical_skills(requirements)

            # Get updated text for verification
            updated_text = '\n'.join([p.text for p in self.doc.paragraphs])

            # Verify keywords (no extra API calls - just text comparison)
            verification = self.verify_keywords_added(original_text, updated_text, missing_skills)

            # Print verification results
            print('\n' + '='*60)
            print('📊 KEYWORD VERIFICATION')
            print('='*60)

            if verification['added']:
                print(f"\n✅ Added ({len(verification['added'])} keywords):")
                for kw in verification['added']:
                    print(f"   • {kw}")

            if verification['already_present']:
                print(f"\n⚪ Already Present ({len(verification['already_present'])} keywords):")
                for kw in verification['already_present'][:5]:  # Show first 5
                    print(f"   • {kw}")
                if len(verification['already_present']) > 5:
                    print(f"   ... and {len(verification['already_present']) - 5} more")

            if verification['not_added']:
                print(f"\n❌ Not Added ({len(verification['not_added'])} keywords):")
                for kw in verification['not_added']:
                    print(f"   • {kw}")

            output_path = self.save_resume()

            print('\n' + '='*60)
            print('✅ UPDATE COMPLETE!')
            print('='*60)

            return output_path, verification['added']

        except Exception as e:
            print(f'\n❌ Error: {e}')
            import traceback
            traceback.print_exc()
            return None, []


def main():
    """Main function"""
    print('='*60)
    print('RESUME UPDATER')
    print('='*60)
    
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_dir = '/Users/gokul/Desktop/Devops 12+/edited resumes'
    
    print(f'Output: {output_dir}\n')
    
    requirements_file = os.path.join(script_dir, 'job_requirement.txt')
    
    if not os.path.exists(requirements_file):
        print('❌ job_requirement.txt not found!')
        sys.exit(1)
    
    print('✓ Found job_requirement.txt')
    
    with open(requirements_file, 'r', encoding='utf-8') as f:
        requirements_text = f.read()
    
    if not requirements_text.strip():
        print('❌ job_requirement.txt is empty!')
        sys.exit(1)
    
    print(f'✓ Loaded requirements')
    
    resume_files = [f for f in os.listdir(script_dir) if f.endswith('.docx') and 'Updated' not in f and '~$' not in f]
    
    if not resume_files:
        print('❌ No resume found!')
        sys.exit(1)
    
    if len(resume_files) == 1:
        resume_file = resume_files[0]
        print(f'✓ Found resume: {resume_file}')
    else:
        print('Multiple resumes found:')
        for i, f in enumerate(resume_files, 1):
            print(f'  {i}. {f}')
        choice = input('\nSelect: ')
        try:
            resume_file = resume_files[int(choice) - 1]
        except (ValueError, IndexError):
            print('Invalid!')
            sys.exit(1)
    
    resume_path = os.path.join(script_dir, resume_file)
    
    try:
        updater = ResumeUpdater(resume_path, output_dir)
        output_path = updater.update_resume(requirements_text)
        
        if output_path:
            print(f'\n📄 Resume: {output_path}')
    
    except Exception as e:
        print(f'\n❌ Error: {e}')
        sys.exit(1)


if __name__ == '__main__':
    main()
