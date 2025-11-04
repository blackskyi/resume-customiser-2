#!/usr/bin/env python3
"""
Local test script for resume customization with Claude API
"""
import os
import sys
from resume_updater import ResumeUpdater

# Test job description with skills that need Claude API
TEST_JOB_DESCRIPTION = """
Senior Site Reliability Engineer (SRE)

We're looking for an experienced SRE to join our infrastructure team.

Required Skills:
- Strong SRE practices (SLOs, error budgets, incident response)
- Observability stack: Grafana, Prometheus, Loki, Tempo
- Cloud migration experience (on-premises to AWS/Azure)
- Kubernetes orchestration and Docker containers
- Infrastructure as Code (Terraform, Pulumi)
- CI/CD pipelines (Jenkins, GitLab CI, ArgoCD)
- Python and Go programming
- Service mesh (Istio or Linkerd)

Responsibilities:
- Design and implement monitoring and observability solutions
- Lead cloud migration initiatives
- Build and maintain CI/CD pipelines
- Implement distributed tracing with Loki and Tempo
- Establish SRE practices and on-call processes
"""

def main():
    print("=" * 70)
    print("RESUME CUSTOMIZER - LOCAL TEST")
    print("=" * 70)

    # Check for API key
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        print("\n❌ ANTHROPIC_API_KEY not set!")
        print("\nTo set it temporarily for this test:")
        print("  export ANTHROPIC_API_KEY='your-api-key-here'")
        print("  python3 test_local.py")
        print("\nOr add to ~/.zshrc for permanent:")
        print("  echo 'export ANTHROPIC_API_KEY=\"your-api-key-here\"' >> ~/.zshrc")
        print("  source ~/.zshrc")
        print("\n⚠️  Get your API key from: https://console.anthropic.com/settings/keys")
        sys.exit(1)

    print(f"\n✅ ANTHROPIC_API_KEY is set (length: {len(api_key)} chars)")

    # Paths
    original_resume = '/Users/gokul/Desktop/Devops 12+/untitled folder/Gokul Prasanna Kumar Senior Devops Engineer.docx'
    output_dir = '/tmp'

    if not os.path.exists(original_resume):
        print(f"\n❌ Resume not found: {original_resume}")
        sys.exit(1)

    print(f"\n📄 Original resume: {original_resume}")
    print(f"📁 Output directory: {output_dir}")

    print("\n" + "=" * 70)
    print("STARTING CUSTOMIZATION TEST")
    print("=" * 70)

    try:
        # Create updater instance
        updater = ResumeUpdater(original_resume, output_dir)

        # Run customization
        output_file, added_keywords = updater.update_resume(TEST_JOB_DESCRIPTION)

        if output_file:
            print("\n" + "=" * 70)
            print("✅ SUCCESS!")
            print("=" * 70)
            print(f"\n📥 Customized resume saved to: {output_file}")

            # Analyze what was added
            from docx import Document
            original_doc = Document(original_resume)
            custom_doc = Document(output_file)

            original_text = '\n'.join([p.text for p in original_doc.paragraphs])
            custom_text = '\n'.join([p.text for p in custom_doc.paragraphs])

            char_diff = len(custom_text) - len(original_text)

            print(f"\n📊 Statistics:")
            print(f"   • Original: {len(original_text)} characters")
            print(f"   • Customized: {len(custom_text)} characters")
            print(f"   • Added: +{char_diff} characters")

            # Check for target skills
            print(f"\n🎯 Checking for target skills:")
            target_skills = ['Loki', 'Tempo', 'Migration', 'SRE', 'Site Reliability']
            for skill in target_skills:
                in_original = skill.lower() in original_text.lower()
                in_custom = skill.lower() in custom_text.lower()

                if in_custom and not in_original:
                    print(f"   ✅ {skill}: ADDED")
                elif in_custom and in_original:
                    print(f"   ⚪ {skill}: Already in original")
                else:
                    print(f"   ❌ {skill}: NOT ADDED")

            print(f"\n✅ Test complete! Check the output file.")
        else:
            print("\n❌ Resume customization failed - no output file generated")
            sys.exit(1)

    except RuntimeError as e:
        print(f"\n❌ CRITICAL ERROR: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == '__main__':
    main()
